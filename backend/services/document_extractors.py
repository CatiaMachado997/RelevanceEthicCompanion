"""
Document text extractors.

extract_text() routes a file path to the appropriate extractor based on
extension or MIME type and returns plain text. Heavy dependencies (pypdf,
python-docx) are imported lazily so cold start is not slowed for users
uploading plain text only.
"""

from __future__ import annotations

import logging
import os
from typing import Optional

logger = logging.getLogger(__name__)


_PDF_MIMES = {"application/pdf"}
_DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}
_TEXT_MIME_PREFIX = "text/"


def _is_pdf(ext: str, mime: str) -> bool:
    return ext == ".pdf" or mime in _PDF_MIMES


def _is_docx(ext: str, mime: str) -> bool:
    return ext == ".docx" or mime in _DOCX_MIMES or "wordprocessingml" in mime


def _is_text(ext: str, mime: str) -> bool:
    if ext in {".txt", ".md", ".markdown"}:
        return True
    if mime.startswith(_TEXT_MIME_PREFIX):
        return True
    # No extension and no mime hint — treat as text (best-effort).
    if not ext and not mime:
        return True
    return False


def extract_text(file_path: str, mime_type: Optional[str] = None) -> str:
    """Return plain text from a file. Routes by extension/mime type.

    Returns empty string for unsupported types (logs a warning) so the
    pipeline does not crash. Image-only PDF pages are skipped with a warning.
    """
    ext = os.path.splitext(file_path)[1].lower()
    mime = (mime_type or "").split(";")[0].strip().lower()

    if _is_pdf(ext, mime):
        return _extract_pdf(file_path)
    if _is_docx(ext, mime):
        return _extract_docx(file_path)
    if _is_text(ext, mime):
        return _extract_text(file_path)

    logger.warning(
        "extract_text: unsupported file type ext=%s mime=%s path=%s",
        ext,
        mime,
        file_path,
    )
    return ""


def _extract_pdf(file_path: str) -> str:
    import pypdf  # lazy import

    try:
        reader = pypdf.PdfReader(file_path)
    except Exception as e:
        logger.warning("PDF open failed for %s: %s", file_path, e)
        return ""

    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as e:
            logger.warning("PDF page %d extract failed in %s: %s", i, file_path, e)
            continue
        if text.strip():
            pages.append(text)
        else:
            logger.warning(
                "PDF page %d in %s yielded no extractable text (image-only?)",
                i,
                file_path,
            )
    return "\n\n".join(pages)


def _extract_docx(file_path: str) -> str:
    import docx  # lazy import — provided by the python-docx package

    try:
        document = docx.Document(file_path)
    except Exception as e:
        logger.warning("DOCX open failed for %s: %s", file_path, e)
        return ""

    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    return "\n".join(parts)


def _extract_text(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception as e:
        logger.warning("Text read failed for %s: %s", file_path, e)
        return ""
