"""Tests for services.document_extractors.extract_text."""

from __future__ import annotations

import logging
import os

import pytest

from services.document_extractors import extract_text


@pytest.fixture
def pdf_fixture(tmp_path):
    """Generate a 1-page PDF containing 'Hello PDF world.' text."""
    from pypdf import PdfWriter

    # pypdf can't create text from scratch, but reportlab is heavy.
    # Use the lower-level approach: build a minimal PDF via reportlab if
    # available, otherwise fall back to pypdf merging an in-memory drawn page.
    try:
        from reportlab.pdfgen import canvas  # type: ignore

        path = tmp_path / "hello.pdf"
        c = canvas.Canvas(str(path))
        c.drawString(100, 750, "Hello PDF world.")
        c.showPage()
        c.save()
        return str(path)
    except ImportError:
        pass

    # Fallback: hand-crafted minimal PDF with a text object.
    # This is a tiny but valid PDF that pypdf can parse and extract text from.
    pdf_bytes = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
        b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"4 0 obj<</Length 56>>stream\n"
        b"BT /F1 12 Tf 100 700 Td (Hello PDF world.) Tj ET\n"
        b"endstream\nendobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"xref\n0 6\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000052 00000 n \n"
        b"0000000101 00000 n \n"
        b"0000000192 00000 n \n"
        b"0000000291 00000 n \n"
        b"trailer<</Size 6/Root 1 0 R>>\n"
        b"startxref\n352\n%%EOF\n"
    )
    path = tmp_path / "hello.pdf"
    path.write_bytes(pdf_bytes)
    # Sanity check that pypdf can open & extract.
    import pypdf

    reader = pypdf.PdfReader(str(path))
    if "Hello PDF world." not in (reader.pages[0].extract_text() or ""):
        # If our hand-rolled PDF doesn't extract cleanly, skip the PDF tests.
        pytest.skip("No PDF generator available (install reportlab to enable)")
    return str(path)


@pytest.fixture
def docx_fixture(tmp_path):
    """Generate a .docx containing 'Hello DOCX world.'."""
    import docx

    path = tmp_path / "hello.docx"
    document = docx.Document()
    document.add_paragraph("Hello DOCX world.")
    document.save(str(path))
    return str(path)


def test_extract_pdf(pdf_fixture):
    text = extract_text(pdf_fixture, mime_type="application/pdf")
    assert "Hello PDF world." in text


def test_extract_pdf_by_extension_only(pdf_fixture):
    # No mime type passed; routes purely by .pdf extension.
    text = extract_text(pdf_fixture)
    assert "Hello PDF world." in text


def test_extract_docx(docx_fixture):
    text = extract_text(
        docx_fixture,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert "Hello DOCX world." in text


def test_extract_docx_by_extension_only(docx_fixture):
    text = extract_text(docx_fixture)
    assert "Hello DOCX world." in text


def test_extract_docx_includes_table_cells(tmp_path):
    import docx

    path = tmp_path / "table.docx"
    document = docx.Document()
    document.add_paragraph("Intro paragraph.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "CellOne"
    table.rows[0].cells[1].text = "CellTwo"
    document.save(str(path))

    text = extract_text(str(path))
    assert "Intro paragraph." in text
    assert "CellOne" in text
    assert "CellTwo" in text


def test_extract_plain_text(tmp_path):
    path = tmp_path / "note.txt"
    path.write_text("Hello plain text world.\nLine two.", encoding="utf-8")
    text = extract_text(str(path), mime_type="text/plain")
    assert text == "Hello plain text world.\nLine two."


def test_extract_markdown(tmp_path):
    path = tmp_path / "note.md"
    path.write_text("# Heading\n\nBody.", encoding="utf-8")
    text = extract_text(str(path))
    assert "# Heading" in text
    assert "Body." in text


def test_unknown_extension_returns_empty_with_warning(tmp_path, caplog):
    path = tmp_path / "blob.xyz"
    path.write_bytes(b"\x00\x01\x02")
    with caplog.at_level(logging.WARNING, logger="services.document_extractors"):
        text = extract_text(str(path), mime_type="application/octet-stream")
    assert text == ""
    assert any("unsupported file type" in rec.message for rec in caplog.records)
